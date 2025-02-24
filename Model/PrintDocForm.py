import os
import subprocess
import sys
import docx
from config_helper import config

class PrintDocForm():
    def __init__(self, partition: str):
        '''partition - имя раздела программмы'''
        path_export = config['Helper_export']['path_export']
        self.path_export_partition = path_export + '/' + partition
        if not os.path.isdir(self.path_export_partition):
            os.mkdir(self.path_export_partition)
        self.path_helper = config['Setting_helper']['path_helper']

    def replace_text(self, forma: str, dictionary: dict, filename: str):
        '''Заменяет текст в указанной форме,
            forma-имя файла из папки Form_print,
            dict-словарь для замены слов в форме,
            filename-имя файла при сохранении'''
        forma = self.path_helper + '/' + 'Form_print' + '/' + forma
        doc = docx.Document(forma)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = docx.shared.Pt(12)

        for i in dictionary:
            for p in doc.paragraphs:
                if p.text.find(i) >= 0:
                    p.text = p.text.replace(i, dictionary[i])

        for j in dictionary:
            for table in doc.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for p in cell.paragraphs:
                            if p.text.find(j) >= 0:
                                p.text = p.text.replace(j, dictionary[j])

        path_save_file=self.path_export_partition+'/'+filename+'.docx'

        doc.save(path_save_file)
        opener = "open" if sys.platform == "darwin" else "xdg-open"
        subprocess.call([opener, path_save_file])
