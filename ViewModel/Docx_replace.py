import subprocess
import sys

import docx
def replace_text(forma, dictionary, saveFileName):
    '''Заменяет текст в указанной форме'''
    doc = docx.Document(forma)
    style=doc.styles['Normal']
    font=style.font
    font.name='Times New Roman'
    font.size=docx.shared.Pt(12)

    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i)>=0:
                p.text=p.text.replace(i, dictionary[i])

    for j in dictionary:
        for table in doc.tables:
            for col in table.columns:
                for cell in col.cells:
                    for p in cell.paragraphs:
                        if p.text.find(j)>=0:
                            p.text=p.text.replace(j, dictionary[j])


    doc.save(saveFileName+'.docx')
    opener = "open" if sys.platform == "darwin" else "xdg-open"
    subprocess.call([opener, saveFileName+'.docx'])


